import tempfile
import unittest
from pathlib import Path

import docx

from app.services.cv_parser import _parse_docx, _parse_pdf_page


class FakePdfPage:
    def __init__(self, words: list[dict[str, object]], width: float = 600.0, height: float = 800.0):
        self._words = words
        self.width = width
        self.height = height

    def extract_words(self, *args, **kwargs):
        return self._words

    def extract_text(self, layout: bool = False):
        return ""


class CvParserTests(unittest.TestCase):
    def test_parse_pdf_page_keeps_column_order(self):
        page = FakePdfPage(
            [
                {"text": "Experience", "x0": 10, "top": 10},
                {"text": "Built", "x0": 10, "top": 24},
                {"text": "product", "x0": 50, "top": 24},
                {"text": "Skills", "x0": 320, "top": 10},
                {"text": "Python", "x0": 320, "top": 24},
                {"text": "SQL", "x0": 380, "top": 24},
            ]
        )

        parsed = _parse_pdf_page(page)

        self.assertEqual(parsed, "Experience\nBuilt product\nSkills\nPython SQL")

    def test_parse_docx_preserves_table_blocks(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            file_path = Path(temp_dir) / "resume.docx"

            document = docx.Document()
            document.add_paragraph("Header")
            table = document.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "Skills"
            table.cell(0, 1).text = "Python"
            document.add_paragraph("Footer")
            document.save(file_path)

            parsed = _parse_docx(str(file_path))

        self.assertIn("Header", parsed)
        self.assertIn("Skills | Python", parsed)
        self.assertIn("Footer", parsed)
        self.assertLess(parsed.index("Header"), parsed.index("Skills | Python"))
        self.assertLess(parsed.index("Skills | Python"), parsed.index("Footer"))


if __name__ == "__main__":
    unittest.main()