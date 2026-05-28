import os
from typing import Any, Iterable

import docx
import pdfplumber
from docx.table import Table
from docx.text.paragraph import Paragraph


def _word_text(word: dict[str, Any]) -> str:
    return str(word.get("text") or "").strip()


def _word_x0(word: dict[str, Any]) -> float:
    return float(word.get("x0") or 0.0)


def _word_top(word: dict[str, Any]) -> float:
    return float(word.get("top") or word.get("doctop") or 0.0)


def _words_to_line(words: list[dict[str, Any]]) -> str:
    ordered = sorted(words, key=_word_x0)
    parts: list[str] = []
    for word in ordered:
        text = _word_text(word)
        if text:
            parts.append(text)
    return " ".join(parts).strip()


def _group_words_into_lines(words: list[dict[str, Any]], y_tolerance: float = 4.0) -> list[tuple[float, float, str]]:
    if not words:
        return []

    ordered = sorted(words, key=lambda word: (_word_top(word), _word_x0(word)))
    lines: list[tuple[float, float, str]] = []

    current_words: list[dict[str, Any]] = []
    current_top: float | None = None

    for word in ordered:
        top = _word_top(word)
        if current_top is None or abs(top - current_top) <= y_tolerance:
            current_words.append(word)
            if current_top is None:
                current_top = top
            else:
                current_top = (current_top + top) / 2
            continue

        line_text = _words_to_line(current_words)
        if line_text:
            lines.append((current_top or 0.0, min(_word_x0(item) for item in current_words), line_text))
        current_words = [word]
        current_top = top

    if current_words:
        line_text = _words_to_line(current_words)
        if line_text:
            lines.append((current_top or 0.0, min(_word_x0(item) for item in current_words), line_text))

    return lines


def extract_pdf_structured(file_path: str) -> list[str]:
    text_blocks: list[str] = []

    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            words = page.extract_words(use_text_flow=True) or []
            lines = _group_words_into_lines([word for word in words if _word_text(word)])

            if not lines:
                page_text = page.extract_text(layout=True) or page.extract_text() or ""
                if page_text.strip():
                    text_blocks.extend(line.strip() for line in page_text.splitlines() if line.strip())
                continue

            lines.sort(key=lambda block: (block[0], block[1]))
            for _, _, line_text in lines:
                if line_text.strip():
                    text_blocks.append(line_text.strip())

    return text_blocks


def _detect_column_boundaries(words: list[dict[str, Any]], page_width: float) -> list[float]:
    """Detect potential column split positions"""
    if not words:
        return []
    
    # Get all x positions
    x_positions = sorted({round(_word_x0(word), 1) for word in words if _word_text(word)})
    
    if len(x_positions) < 4:
        return []
    
    # Find gaps between clusters of x positions
    gaps = []
    for i in range(len(x_positions) - 1):
        gap = x_positions[i + 1] - x_positions[i]
        gaps.append((gap, i))
    
    if not gaps:
        return []
    
    # Find the largest gap
    max_gap, max_gap_idx = max(gaps, key=lambda x: x[0])
    
    # Check if this is a meaningful column gap (at least 10% of page width)
    if max_gap < page_width * 0.08:  # 8% of page width
        return []
    
    # Get the split position
    split_x = (x_positions[max_gap_idx] + x_positions[max_gap_idx + 1]) / 2
    
    # Verify that both sides have reasonable content
    left_words = [w for w in words if _word_x0(w) < split_x and _word_text(w)]
    right_words = [w for w in words if _word_x0(w) > split_x and _word_text(w)]
    
    if len(left_words) < 3 or len(right_words) < 3:
        return []
    
    return [split_x]


def _split_words_into_regions(words: list[dict[str, Any]], page_width: float) -> list[tuple[float, float, list[dict[str, Any]]]]:
    """Split words into column regions based on x-position clustering"""
    boundaries = _detect_column_boundaries(words, page_width)
    
    if not boundaries:
        return [(0.0, page_width, words)]
    
    # Create regions based on boundaries
    regions: list[tuple[float, float]] = []
    left = 0.0
    for boundary in sorted(boundaries):
        regions.append((left, boundary))
        left = boundary
    regions.append((left, page_width))
    
    # Assign words to regions
    assigned_regions = []
    for region_left, region_right in regions:
        region_words = []
        for word in words:
            x0 = _word_x0(word)
            if region_left <= x0 < region_right and _word_text(word):
                region_words.append(word)
        
        # Only include regions with substantial content
        if len(region_words) >= 3:  # Minimum words per column
            assigned_regions.append((region_left, region_right, region_words))
    
    # If we have at least 2 good regions, use them
    if len(assigned_regions) >= 2:
        return assigned_regions
    
    # Fallback to single column
    return [(0.0, page_width, words)]


def _extract_pdf_page_text(page) -> str:
    """Extract text with proper layout handling."""
    words = [word for word in page.extract_words(use_text_flow=True) if _word_text(word)]
    if not words:
        return page.extract_text(layout=True) or page.extract_text() or ""
    
    regions = _split_words_into_regions(words, float(page.width))
    
    if len(regions) <= 1:
        # Single column - simple extraction
        lines = []
        for top, left, line_text in _group_words_into_lines(words):
            if line_text:
                lines.append((top, left, line_text))
        lines.sort(key=lambda block: (block[0], block[1]))  # Sort by top, then left
        return "\n".join(block[2] for block in lines)
    
    # Multiple columns - process each column separately
    column_texts = []
    
    for region_left, region_right, region_words in regions:
        # Group words in this column into lines
        lines = _group_words_into_lines(region_words)
        
        # Sort lines by vertical position (top to bottom)
        lines.sort(key=lambda block: block[0])
        
        # Extract just the text from lines
        column_text = "\n".join(block[2] for block in lines if block[2])
        if column_text.strip():
            column_texts.append(column_text)
    
    # Keep column order linear so downstream chunking sees a stable reading order.
    return "\n".join(column_texts)


def _parse_pdf_page(page) -> str:
    return _extract_pdf_page_text(page)


def _paragraph_text(paragraph: Paragraph) -> str:
    return paragraph.text.strip()


def _table_text(table: Table) -> str:
    rows: list[str] = []
    for row in table.rows:
        cells: list[str] = []
        for cell in row.cells:
            cell_text = "\n".join(
                _paragraph_text(paragraph)
                for paragraph in cell.paragraphs
                if _paragraph_text(paragraph)
            ).strip()
            if cell_text:
                cells.append(cell_text)
        if cells:
            rows.append(" | ".join(cells))

    return "\n".join(rows).strip()


def _iter_docx_blocks(document: docx.Document) -> Iterable[str]:
    for child in document.element.body.iterchildren():
        if child.tag.endswith("}p"):
            paragraph = Paragraph(child, document)
            paragraph_text = _paragraph_text(paragraph)
            if paragraph_text:
                yield paragraph_text
        elif child.tag.endswith("}tbl"):
            table = Table(child, document)
            table_text = _table_text(table)
            if table_text:
                yield table_text


def _parse_pdf(file_path: str) -> str:
    return "\n".join(extract_pdf_structured(file_path))


def _parse_docx(file_path: str) -> str:
    document = docx.Document(file_path)
    blocks = [block for block in _iter_docx_blocks(document) if block.strip()]
    if blocks:
        return "\n".join(blocks)

    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


def parse_file(file_path: str, filename: str) -> str:
    _, ext = os.path.splitext(filename.lower())
    if ext == ".pdf":
        return _parse_pdf(file_path)
    if ext == ".docx":
        return _parse_docx(file_path)

    raise ValueError("Only PDF and DOCX files are supported")
